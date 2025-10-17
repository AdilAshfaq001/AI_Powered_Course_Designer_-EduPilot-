import streamlit as st
import asyncio
import os
import io
from docx import Document
from modules.module3_content_generator import generate_course_content

def local_css(file_name):
    try:
        with open(file_name) as f:
            st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)
    except FileNotFoundError:
        st.error(f"CSS file not found: {file_name}. Make sure the 'styles' directory and 'style.css' file exist.")

# --- Apply the custom CSS ---
local_css("HtmlCSS\streamlit.css")

st.set_page_config(layout="wide")
st.title("EduPilot — Module 3: Weekly Content Generator")

# --- Input Section in Sidebar ---
with st.sidebar:
    st.header("⚙️ Content Parameters")
    uploaded_file = st.file_uploader("Upload Module 2 JSON file", type=["json"])

    week_number = st.slider(
        "Select Week to Generate Content For",
        min_value=1,
        max_value=15,
        value=1,
        step=1
    )

    complexity = st.selectbox(
        "Select Learning Complexity",
        options=["Beginner", "Intermediate", "Advanced", "Expert"],
        index=1
    )

    multimedia_prefs = st.multiselect(
        "Select Multimedia/Resource Preferences",
        ['Text', 'Books', 'Articles', 'Websites', 'Official Documentation', 'YouTube Links'],
        default=['Text', 'Books', 'Articles', 'YouTube Links']
    )

    submitted = st.button("Generate Weekly Content")


# --- Output Section in Main Area ---
if submitted:
    if uploaded_file:
        temp_dir = "temp_uploads"
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, uploaded_file.name)
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        with st.spinner(f"🚀 Generating content for Week {week_number}..."):
            content, filepath = asyncio.run(generate_course_content(
                temp_path, week_number, complexity, multimedia_prefs
            ))
            
            if content:
                st.success(f"✅ Content generation for Week {week_number} complete!")
                st.subheader(f"Generated Content for Week {week_number}")

                # Display content in expanders
                with st.expander("📝 Lecture Notes"):
                    st.markdown(content.get("lecture_notes", "Not available."))

                with st.expander("📚 Reading Materials & References"):
                    st.markdown(content.get("reading_materials", "Not available."))

                with st.expander("💻 Exercises & Projects"):
                    st.markdown(content.get("exercises_projects", "Not available."))

                with st.expander("❓ Assessment Questions"):
                    st.markdown(content.get("assessment_questions", "Not available."))

                # --- Download Buttons ---
                st.subheader("⬇️ Download Options")
                col1, col2 = st.columns(2)

                # JSON download
                with col1:
                    with open(filepath, "r", encoding="utf-8") as f:
                        st.download_button(
                            label=f"📥 Download as JSON",
                            data=f.read(),
                            file_name=os.path.basename(filepath),
                            mime="application/json",
                            use_container_width=True
                        )

                # DOCX download
                with col2:
                    doc = Document()
                    doc.add_heading(f'Course Content: Week {week_number}', level=1)
                    doc.add_heading('Lecture Notes', level=2)
                    doc.add_paragraph(content.get("lecture_notes", "Not available."))
                    doc.add_heading('Reading Materials & References', level=2)
                    doc.add_paragraph(content.get("reading_materials", "Not available."))
                    doc.add_heading('Exercises & Projects', level=2)
                    doc.add_paragraph(content.get("exercises_projects", "Not available."))
                    doc.add_heading('Assessment Questions', level=2)
                    doc.add_paragraph(content.get("assessment_questions", "Not available."))

                    bio = io.BytesIO()
                    doc.save(bio)
                    bio.seek(0)
                    
                    st.download_button(
                        label="📥 Download as DOCX",
                        data=bio.getvalue(),
                        file_name=f"Week_{week_number}_Content.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            else:
                st.error(f"⚠️  Failed to generate content for Week {week_number}.")

    else:
        st.error("Please upload the JSON file from Module 2.")
else:
    st.info("Upload the curriculum file and set the parameters in the sidebar, then click 'Generate Weekly Content'.")

