import streamlit as st
from modules.module2_curriculum_structurer import generate_curriculum
import os
import re
import io
from docx import Document

def local_css(file_name):
    try:
        with open(file_name) as f:
            st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)
    except FileNotFoundError:
        st.error(f"CSS file not found: {file_name}. Make sure the 'styles' directory and 'style.css' file exist.")

# --- Apply the custom CSS ---
local_css("HtmlCSS\streamlit.css")

st.set_page_config(layout="wide")
st.title("EduPilot — Module 2: Curriculum Structurer")

# --- Input Section in Sidebar ---
with st.sidebar:
    st.header("⚙️ Curriculum Parameters")
    uploaded_file = st.file_uploader("Upload Module 1 JSON file", type=["json"])
    learning_outcomes_input = st.text_area("Or paste Learning Outcomes here", height=150)
    time_weeks = st.slider("Course Duration (weeks)", 8, 20, 15)
    approach = st.selectbox("Pedagogical Approach", ["Project-based", "Theory", "Blended"])
    assessments = st.multiselect("Assessment Preferences", ["Quizzes", "Projects", "Exams", "Presentations", "Labs"], default=["Quizzes", "Projects", "Exams"])
    submitted = st.button("Generate Curriculum")

# --- Output Section in Main Area ---
if submitted:
    if not uploaded_file and not learning_outcomes_input:
        st.error("Please upload a file or paste learning outcomes to proceed.")
    else:
        if uploaded_file:
            temp_dir = "temp_uploads"
            os.makedirs(temp_dir, exist_ok=True)
            temp_path = os.path.join(temp_dir, uploaded_file.name)
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            learning_input = temp_path
        else:
            learning_input = learning_outcomes_input

        with st.spinner("Generating curriculum..."):
            curriculum_text, _ = generate_curriculum(learning_input, time_weeks, approach, assessments)
            
            if curriculum_text:
                st.success("✅ Curriculum Generated Successfully!")
                st.subheader("📄 Generated Curriculum Plan")

                modules = re.split(r"(?=\*\*Module \d+)", curriculum_text)
                for mod in modules:
                    if mod.strip():
                        header_match = re.match(r"\*\*(Module \d+.*?)\*\*", mod)
                        header = header_match.group(1) if header_match else "Module"
                        body = mod.replace(f"**{header}**", "").strip()
                        with st.expander(header):
                            st.markdown(body, unsafe_allow_html=True)

                # --- Download Buttons ---
                st.subheader("⬇️ Download Options")
                col1, col2 = st.columns(2)

                with col1:
                    st.download_button(
                        label="📥 Download as TXT",
                        data=curriculum_text,
                        file_name="Curriculum_Plan.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                
                with col2:
                    doc = Document()
                    doc.add_heading("Curriculum Plan", level=1)
                    for mod in modules:
                        if mod.strip():
                            header_match = re.match(r"\*\*(Module \d+.*?)\*\*", mod)
                            header = header_match.group(1) if header_match else "Module"
                            body = mod.replace(f"**{header}**", "").strip()
                            doc.add_paragraph(header, style='Heading2')
                            doc.add_paragraph(body)
                    
                    bio = io.BytesIO()
                    doc.save(bio)
                    bio.seek(0)
                    
                    st.download_button(
                        label="📥 Download as DOCX",
                        data=bio.getvalue(),
                        file_name="Curriculum_Plan.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            else:
                st.error("⚠️ Failed to generate curriculum.")
else:
    st.info("Provide the learning outcomes and course parameters in the sidebar, then click 'Generate Curriculum'.")
